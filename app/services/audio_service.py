import os
import whisper
from docx import Document

def audio_to_text(path_to_file, filename):
    model = whisper.load_model('base')
    result = model.transcribe(path_to_file, fp16=False)
    doc = Document()
    
    p = doc.add_paragraph()
    title = p.add_run(f'Transcripción de {filename}')
    title.bold = True

    body = doc.add_paragraph(result['text'])
    doc_name = f"{path_to_file}.docx"
    doc.save(f"{path_to_file}.docx")
    path_to_doc = f"{path_to_file}.docx"
    print(path_to_doc)
    
    return {"path": path_to_doc, "filename": f"{filename[:-4]}.docx"}


def upload_file(file):
    temp_dir = "temp"
    os.makedirs(temp_dir, exist_ok=True)
    temp_file_path = os.path.join(temp_dir, file.filename)

    with open(temp_file_path, "wb") as temp_file:
        temp_file.write(file.file.read())
        
    transcription = audio_to_text(temp_file_path, file.filename)
    
    return transcription
        
        